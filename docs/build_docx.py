# -*- coding: utf-8 -*-
"""Build both Word/PDF deliverables from docs/GPTEnergy.html via the html2doc skill.

  JoulePoint_1col.docx/pdf  -- PRIMARY: single-column, horizontal multi-panel figures,
                              figures height-capped so figure+caption bumps less (less white).
  JoulePoint_2col.docx/pdf  -- best-effort: two-column, multi-panel figures swapped to their
                              VERTICAL stacked variants so they flow single-column (no float gaps).

Run from docs/:  python build_docx.py
Word (win32com) is used to render PDFs; on non-Windows, the .docx are still produced.
"""
import os, shutil, subprocess, sys

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL = r"C:\Users\apart\.claude\skills\html2doc"
PY = sys.executable
NODE_ENV = {**os.environ, "NODE_PATH": os.path.join(SKILL, "node_modules")}


def sh(cmd, env=None):
    print("+", " ".join(cmd))
    subprocess.run(cmd, cwd=HERE, env=env, check=True)


def stage1(src, out):
    sh(["node", os.path.join(SKILL, "scripts", "katex_to_mathml.js"), "--input", src, "--output", out], env=NODE_ENV)


def stage2(src, out, profile):
    sh([PY, os.path.join(SKILL, "scripts", "convert_to_docx.py"), "--input", src, "--output", out, "--profile", profile])


def stage3(src, out, profile, extra):
    sh([PY, os.path.join(SKILL, "scripts", "apply_academic_style.py"), "--input", src, "--output", out,
        "--profile", profile, "--font-family", "Georgia"] + extra)


NAVY, SOFT, MUTED = (0x14, 0x38, 0x5C), (0x2C, 0x31, 0x38), (0x5A, 0x62, 0x6C)


def polish_premium(path):
    """House-palette pass on a built DOCX: navy title/headings/caption labels,
    small-caps abstract label, muted affiliations and caption bodies, and the
    key-finding paragraph as a navy-ruled shaded callout box."""
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(path)

    BLACK = (0x00, 0x00, 0x00)

    def tint(runs, rgb):
        for r in runs:
            r.font.color.rgb = RGBColor(*BLACK)

    for p in doc.paragraphs:
        name = p.style.name
        if name == "Title":
            tint(p.runs, NAVY)
        elif name.startswith("Heading") or name == "References Heading":
            tint(p.runs, NAVY)
        elif name == "Abstract Label":
            tint(p.runs, NAVY)
            for r in p.runs:
                r.font.small_caps = True
        elif name in ("Image Caption", "Table Caption"):
            for r in p.runs:
                if r.bold:
                    r.font.color.rgb = RGBColor(*BLACK)
                else:
                    r.font.color.rgb = RGBColor(*BLACK)
        if p.text.startswith("Under load, one static cap per card"):
            ppr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "6"); left.set(qn("w:color"), "14385C")
            borders.append(left); ppr.append(borders)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F5F8")
            ppr.append(shd)
            p.paragraph_format.left_indent = Pt(6)
            for r in p.runs:
                r.font.color.rgb = RGBColor(*BLACK)
                r.font.italic = True

    # author/affiliation front matter lives in a borderless table
    for t in doc.tables:
        cell_text = " ".join(c.text for row in t.rows for c in row.cells)
        if "Apartsin" in cell_text and "Holon" in cell_text:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(*BLACK)
    # tighten references so the list does not spill two entries onto a final page
    from docx.shared import Pt as _Pt
    try:
        doc.styles["Reference Entry"].font.size = _Pt(8.5)
    except KeyError:
        pass

    # balance the final page's columns: end the last columned section with a
    # continuous break so Word levels the reference list across both columns
    from docx.enum.section import WD_SECTION_START
    doc.add_section(WD_SECTION_START.CONTINUOUS)

    for attempt in (1, 2):
        try:
            doc.save(path); break
        except MemoryError:
            if attempt == 2:
                raise
    print("  polished (premium palette)", path)


def render_pdf(docx, pdf):
    try:
        import win32com.client as w, pythoncom
        pythoncom.CoInitialize()
        app = w.DispatchEx("Word.Application"); app.Visible = False
        d = app.Documents.Open(os.path.join(HERE, docx), False, True)
        d.ExportAsFixedFormat(os.path.join(HERE, pdf), 17); d.Close(False); app.Quit()
        print("  rendered", pdf)
    except Exception as e:
        print("  [skip PDF render]", type(e).__name__, str(e)[:80])


def main():
    # ---- 1-column (primary): horizontal figures, capped height to reduce page-bottom white ----
    stage1("GPTEnergy.html", "_gpte_1col_mathml.html")
    stage2("_gpte_1col_mathml.html", "JoulePoint_1col_conv.docx", "camera-ready-generic")
    stage3("JoulePoint_1col_conv.docx", "JoulePoint_1col.docx", "camera-ready-generic",
           ["--figure-max-height-in", "3.5"])
    render_pdf("JoulePoint_1col.docx", "JoulePoint_1col.pdf")

    # ---- 2-column (best-effort) ----
    # Layout choices for continuous columns (no half-empty columns):
    #  * fig_law_fit and fig_frontier stay HORIZONTAL: wide+short, they span both
    #    columns as full-width floats and text flows on (their stacked variants are
    #    taller than a column's free space, bump to the next column, and leave white).
    #  * fig_energy_u (Figure 3) moves three paragraphs later so the text after it
    #    backfills page 5's right column instead of the figure bumping early.
    import re
    html = open(os.path.join(HERE, "GPTEnergy.html"), encoding="utf-8").read()
    html = html.replace("fig_frontier.png", "fig_frontier_stacked.png")
    # Relocate in-column figures so the text after each one backfills the column it
    # would otherwise leave half-empty (2-col layout only; captions keep the numbers).
    MOVES = [
        ("fig_energy_u",             "(Table&nbsp;1).</p>"),
        ("fig_frontier_stacked",     "</ul>"),
        ("fig_power_budget_stacked", "quantitative analysis.</p>"),
    ]
    for png, marker in MOVES:
        m = re.search(r'<figure><img [^>]*' + png + r'\.png.*?</figure>', html, re.S)
        assert m and html.count(marker) == 1, f"relocation anchor changed: {png}"
        fig = m.group(0)
        html = html.replace(fig, "").replace(marker, marker + "\n" + fig)
    open(os.path.join(HERE, "GPTEnergy_2col_src.html"), "w", encoding="utf-8").write(html)
    stage1("GPTEnergy_2col_src.html", "_gpte_2col_mathml.html")
    stage2("_gpte_2col_mathml.html", "JoulePoint_2col_conv.docx", "two-column")
    stage3("JoulePoint_2col_conv.docx", "JoulePoint_2col.docx", "two-column",
           ["--max-span-height-frac", "0.32", "--figure-max-height-in", "3.0"])
    polish_premium(os.path.join(HERE, "JoulePoint_2col.docx"))
    render_pdf("JoulePoint_2col.docx", "JoulePoint_2col.pdf")

    # content canary
    from docx import Document
    import zipfile
    for f in ("JoulePoint_1col.docx", "JoulePoint_2col.docx"):
        z = zipfile.ZipFile(os.path.join(HERE, f))
        media = len([n for n in z.namelist() if n.startswith("word/media/")])
        print(f"  {f}: {media} figures, {len(Document(os.path.join(HERE, f)).tables)} tables")
        assert media == 6, f"{f}: expected 6 figures, got {media}"
    print("done: 1col (primary) + 2col built, 6/6 figures each.")


if __name__ == "__main__":
    main()
